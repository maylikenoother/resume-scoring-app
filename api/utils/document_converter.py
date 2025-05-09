import os
import tempfile
from typing import Optional
import logging

try:
    from pdfminer.high_level import extract_text as extract_text_from_pdf
    has_pdfminer = True
except ImportError:
    has_pdfminer = False

try:
    import docx
    has_docx = True
except ImportError:
    has_docx = False

logger = logging.getLogger(__name__)

def extract_text_from_docx(docx_bytes):
    """Extract text from a .docx file content"""
    if not has_docx:
        logger.error("python-docx library not installed. Cannot extract text from DOCX files.")
        return None
        
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(docx_bytes)
            temp_file_path = temp_file.name
        
        doc = docx.Document(temp_file_path)
        full_text = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        os.unlink(temp_file_path)
        
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        if 'temp_file_path' in locals():
            try:
                os.unlink(temp_file_path)
            except:
                pass
        return None

def convert_to_text(file_content: bytes, content_type: str) -> Optional[str]:
    """
    Convert various document formats to plain text
    
    Args:
        file_content: The binary content of the file
        content_type: The MIME type of the file
    
    Returns:
        Extracted text content or None if conversion failed
    """
    try:
        if content_type == 'application/pdf':
            if not has_pdfminer:
                logger.warning("pdfminer.six library not installed. Cannot extract text from PDF files.")
                return file_content.decode('utf-8', errors='replace')
                
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                extracted_text = extract_text_from_pdf(temp_file_path)
                return extracted_text if extracted_text.strip() else None
            finally:
                os.unlink(temp_file_path)
            
        elif content_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]:
            extracted_text = extract_text_from_docx(file_content) 
            if extracted_text:
                return extracted_text
            
            return file_content.decode('utf-8', errors='replace')
            
        elif content_type == 'text/plain':
            return file_content.decode('utf-8', errors='replace')
            
        else:
            logger.warning(f"Unsupported file type for text extraction: {content_type}")
            return file_content.decode('utf-8', errors='replace')
            
    except Exception as e:
        logger.error(f"Document conversion error: {str(e)}")
        try:
            return file_content.decode('utf-8', errors='replace')
        except:
            return None